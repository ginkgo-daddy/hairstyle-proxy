import http.client
import json
import os
import mimetypes
from codecs import encode
import time
import requests
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image, ExifTags
import io
import random
import concurrent.futures
import threading
from queue import Queue
import base64
import asyncio
import hashlib
from openai import AsyncOpenAI
from dotenv import load_dotenv
load_dotenv()


def ensure_data_directory():
    """确保数据目录存在并有适当的权限"""
    data_dir = os.environ.get('RAILWAY_VOLUME_MOUNT_PATH', '/data')
    try:
        if not os.path.exists(data_dir):
            os.makedirs(data_dir, exist_ok=True)
            print(f"创建数据目录: {data_dir}")
        
        # 检查目录权限
        if not os.access(data_dir, os.W_OK):
            print(f"警告: 数据目录 {data_dir} 没有写权限")
        else:
            print(f"数据目录就绪: {data_dir}")
            
        return data_dir
    except Exception as e:
        print(f"初始化数据目录失败: {e}")
        # 回退到当前目录
        fallback_dir = os.path.join(os.getcwd(), 'data')
        os.makedirs(fallback_dir, exist_ok=True)
        print(f"使用回退数据目录: {fallback_dir}")
        return fallback_dir

class HairstyleProcessor:
    def __init__(self, api_key=None, webapp_id=None, color_webapp_id=None, max_workers=30, task_timeout=600):
        # 首先确保数据目录存在
        self.data_dir = ensure_data_directory()

        # 从环境变量获取API密钥，如果没有则使用传入的参数
        self.api_key = api_key or os.environ.get('RUNNINGHUB_API_KEY')
        if not self.api_key:
            raise ValueError("API key is required. Set RUNNINGHUB_API_KEY environment variable or pass api_key parameter.")

        # 从环境变量获取Webapp ID，如果没有则使用传入的参数
        self.webapp_id = webapp_id or os.environ.get('RUNNINGHUB_WEBAPP_ID')

        # 从环境变量获取颜色换装Webapp ID
        self.color_webapp_id = color_webapp_id or os.environ.get('RUNNINGHUB_COLOR_WEBAPP_ID')

        # 从环境变量获取OpenRouter API密钥（用于Gemini预处理）
        self.openrouter_api_key = os.environ.get('OPENROUTER_API_KEY')

        self.host = "www.runninghub.cn"
        self.results = []
        self.results_lock = threading.Lock()
        self.max_workers = max_workers
        self.task_timeout = task_timeout  # 每个任务的超时时间（秒），默认600秒

        # 添加时间统计变量
        self.task_times = []  # 存储每次run_hairstyle_task的运行时间
        self.task_count = 0   # 任务总数统计

        # Gemini预处理统计
        self.gemini_times = []  # 存储Gemini预处理时间
        self.gemini_success_count = 0  # 成功预处理数量
        self.gemini_fail_count = 0     # 失败预处理数量

        # 超时统计
        self.timeout_count = 0  # 超时任务数量

    def encode_image(self, image_path):
        """将图像编码为base64字符串，自动处理EXIF方向"""
        try:
            # 使用PIL打开图像并自动处理EXIF方向
            with Image.open(image_path) as img:
                # 自动根据EXIF方向信息旋转图像
                img = self.fix_image_orientation(img)

                # 转换为RGB模式（避免PNG保存问题）
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                # 保存到内存缓冲区
                buffer = io.BytesIO()
                img.save(buffer, format='JPEG', quality=95)
                buffer.seek(0)

                # 编码为base64
                return base64.b64encode(buffer.getvalue()).decode('utf-8')

        except Exception as e:
            print(f"处理图像EXIF方向失败，使用原始方法: {e}")
            # 回退到原始方法
            with open(image_path, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode('utf-8')

    def fix_image_orientation(self, img):
        """根据EXIF信息修正图像方向"""
        try:
            # 使用PIL的ImageOps.exif_transpose方法，这是处理EXIF方向的推荐方法
            from PIL import ImageOps
            img = ImageOps.exif_transpose(img)
            return img
        except ImportError:
            # 如果ImageOps不可用，使用传统方法
            try:
                exif = img._getexif()
                if exif is not None:
                    # 查找方向标签
                    for tag, value in exif.items():
                        if ExifTags.TAGS.get(tag) == 'Orientation':
                            # 根据EXIF方向值旋转图像
                            if value == 2:
                                img = img.transpose(Image.FLIP_LEFT_RIGHT)
                            elif value == 3:
                                img = img.rotate(180, expand=True)
                            elif value == 4:
                                img = img.transpose(Image.FLIP_TOP_BOTTOM)
                            elif value == 5:
                                img = img.transpose(Image.FLIP_LEFT_RIGHT).rotate(90, expand=True)
                            elif value == 6:
                                img = img.rotate(-90, expand=True)
                            elif value == 7:
                                img = img.transpose(Image.FLIP_LEFT_RIGHT).rotate(-90, expand=True)
                            elif value == 8:
                                img = img.rotate(90, expand=True)
                            break
            except Exception as e:
                print(f"修正图像方向失败: {e}")

        return img

    def get_file_hash(self, file_path):
        """计算文件的MD5哈希值"""
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            print(f"计算文件哈希失败: {e}")
            return None

    def save_image_from_base64(self, base64_str, original_path, image_type, file_hash):
        """从base64字符串还原图片并保存，使用哈希值命名"""
        try:
            output_dir = os.path.join(self.data_dir, f"gemini_processed_{image_type}")
            if not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)

            # 使用文件哈希值作为主要标识符，保留原文件名用于识别
            original_filename = os.path.basename(original_path)
            name_without_ext = os.path.splitext(original_filename)[0]

            # 文件名格式: 原名_哈希前8位_gemini_processed.png
            new_filename = f"{name_without_ext}_{file_hash[:8]}_gemini_processed.png"
            filepath = os.path.join(output_dir, new_filename)

            image_data = base64.b64decode(base64_str)
            with open(filepath, "wb") as f:
                f.write(image_data)

            # 创建缓存索引文件
            self.update_cache_index(original_path, filepath, file_hash, image_type)

            return filepath
        except Exception as e:
            print(f"保存图片时出错: {e}")
            return None

    def update_cache_index(self, original_path, processed_path, file_hash, image_type):
        """更新缓存索引文件"""
        try:
            cache_dir = os.path.join(self.data_dir, f"gemini_processed_{image_type}")
            cache_index_path = os.path.join(cache_dir, "cache_index.json")

            # 读取现有索引
            cache_index = {}
            if os.path.exists(cache_index_path):
                try:
                    with open(cache_index_path, 'r', encoding='utf-8') as f:
                        cache_index = json.load(f)
                except:
                    cache_index = {}

            # 更新索引
            cache_index[file_hash] = {
                "original_path": original_path,
                "processed_path": processed_path,
                "timestamp": datetime.now().isoformat(),
                "original_filename": os.path.basename(original_path)
            }

            # 保存索引
            with open(cache_index_path, 'w', encoding='utf-8') as f:
                json.dump(cache_index, f, ensure_ascii=False, indent=2)

        except Exception as e:
            print(f"更新缓存索引失败: {e}")

    def get_cached_processed_path(self, original_path, image_type):
        """基于文件哈希检查是否已有缓存的预处理图片"""
        try:
            # 计算原文件哈希
            file_hash = self.get_file_hash(original_path)
            if not file_hash:
                return None

            # 检查缓存索引
            cache_dir = os.path.join(self.data_dir, f"gemini_processed_{image_type}")
            cache_index_path = os.path.join(cache_dir, "cache_index.json")

            if not os.path.exists(cache_index_path):
                return None

            # 读取缓存索引
            try:
                with open(cache_index_path, 'r', encoding='utf-8') as f:
                    cache_index = json.load(f)
            except:
                return None

            # 查找匹配的哈希
            if file_hash in cache_index:
                cached_info = cache_index[file_hash]
                cached_path = cached_info["processed_path"]

                # 验证缓存文件是否仍然存在
                if os.path.exists(cached_path):
                    return cached_path
                else:
                    # 缓存文件不存在，清理索引
                    del cache_index[file_hash]
                    with open(cache_index_path, 'w', encoding='utf-8') as f:
                        json.dump(cache_index, f, ensure_ascii=False, indent=2)

            return None

        except Exception as e:
            print(f"检查缓存失败: {e}")
            return None

    async def preprocess_image_with_gemini(self, image_path, image_type="user"):
        """使用Gemini对图像进行预处理（异步版本）"""
        thread_name = threading.current_thread().name
        start_time = time.time()

        try:
            print(f"[{thread_name}] 开始Gemini预处理{image_type}图像: {os.path.basename(image_path)}")

            # 检查缓存（基于文件哈希）
            cached_path = self.get_cached_processed_path(image_path, image_type)
            if cached_path:
                print(f"[{thread_name}] ✓ 找到缓存的{image_type}图像: {os.path.basename(cached_path)}")
                return cached_path

            if not self.openrouter_api_key:
                print(f"[{thread_name}] 未设置OPENROUTER_API_KEY，跳过Gemini预处理")
                self.gemini_fail_count += 1
                return image_path

            # 计算文件哈希（用于保存时的文件命名）
            file_hash = self.get_file_hash(image_path)
            if not file_hash:
                print(f"[{thread_name}] 无法计算文件哈希，跳过预处理")
                self.gemini_fail_count += 1
                return image_path

            base64_image = self.encode_image(image_path)

            async with AsyncOpenAI(
                base_url="https://openrouter.ai/api/v1",
                api_key=self.openrouter_api_key,
            ) as client:

                # 根据图片类型设置不同的提示语
                if image_type == "user":
                    prompt_text = "保持人物一致性，保持服饰和发型不变，身材不要太胖，改为半身证件照，光线充足，露出黑色腰带。"
                elif image_type == "hairstyle":
                    prompt_text = "保持人物一致性，保持服饰和发型发色不变，保持发型纹理清晰，光照条件与原图一致，改为半身证件照，露出黑色腰带。"
                else:
                    prompt_text = "保持人物一致性，保持服饰和发型发色不变，保持发型纹理清晰，光照条件与原图一致，改为半身证件照，露出黑色腰带。"

                completion = await client.chat.completions.create(
                    model="google/gemini-2.5-flash-image-preview",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": prompt_text
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{base64_image}"
                                    }
                                }
                            ]
                        }
                    ]
                )

                end_time = time.time()
                elapsed = end_time - start_time
                self.gemini_times.append(elapsed)
                print(f"[{thread_name}] Gemini预处理{image_type}耗时: {elapsed:.2f}秒")

                return await self.process_gemini_response(completion, image_path, image_type, file_hash, thread_name, client, prompt_text, base64_image, attempt=1)

        except Exception as e:
            end_time = time.time()
            elapsed = end_time - start_time
            self.gemini_times.append(elapsed)
            print(f"[{thread_name}] Gemini预处理出错: {e}")
            print(f"[{thread_name}] 使用原图继续处理...")
            self.gemini_fail_count += 1
            return image_path

    async def process_gemini_response(self, completion, image_path, image_type, file_hash, thread_name, client, prompt_text, base64_image, attempt=1):
        """处理Gemini API响应，包含重试机制"""
        max_retries = 2  # 最多重试1次，总共2次尝试

        # 检查响应中是否有图片数据
        if hasattr(completion.choices[0].message, 'images') and completion.choices[0].message.images:
            image_url = completion.choices[0].message.images[0]["image_url"]['url']

            if image_url.startswith("data:image/"):
                base64_data = image_url.split(",")[1]
                processed_image_path = self.save_image_from_base64(
                    base64_data,
                    image_path,    # 原始路径
                    image_type,    # 图像类型
                    file_hash      # 文件哈希
                )

                if processed_image_path:
                    print(f"[{thread_name}] ✓ Gemini{image_type}预处理成功: {os.path.basename(processed_image_path)}")
                    self.gemini_success_count += 1
                    return processed_image_path
                else:
                    print(f"[{thread_name}] 保存失败，使用原图")
                    self.gemini_fail_count += 1
                    return image_path
            else:
                print(f"[{thread_name}] 非base64格式URL，使用原图")
                self.gemini_fail_count += 1
                return image_path
        else:
            # 响应中无图片数据，尝试重试
            if attempt < max_retries:
                print(f"[{thread_name}] 响应中无图片数据，进行第{attempt + 1}次尝试...")
                try:
                    # 等待一小段时间再重试
                    await asyncio.sleep(1)

                    # 重新调用API
                    retry_completion = await client.chat.completions.create(
                        model="google/gemini-2.5-flash-image-preview",
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": prompt_text
                                    },
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:image/jpeg;base64,{base64_image}"
                                        }
                                    }
                                ]
                            }
                        ]
                    )

                    print(f"[{thread_name}] 重试请求完成，处理响应...")
                    # 递归调用处理重试的响应
                    return await self.process_gemini_response(
                        retry_completion, image_path, image_type, file_hash,
                        thread_name, client, prompt_text, base64_image, attempt + 1
                    )

                except Exception as retry_error:
                    print(f"[{thread_name}] 重试请求失败: {retry_error}")
                    print(f"[{thread_name}] 使用原图")
                    self.gemini_fail_count += 1
                    return image_path
            else:
                print(f"[{thread_name}] 达到最大重试次数，响应中仍无图片数据，使用原图")
                self.gemini_fail_count += 1
                return image_path

    def preprocess_images_concurrently(self, user_image_path, hairstyle_image_path):
        """并发预处理用户图片和发型图片（同步接口）"""
        thread_name = threading.current_thread().name
        try:
            print(f"[{thread_name}] 开始并发预处理图像...")

            # 创建新的事件循环来运行异步代码
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

            try:
                # 并发执行两个预处理任务
                processed_user_image, processed_hairstyle_image = loop.run_until_complete(
                    asyncio.gather(
                        self.preprocess_image_with_gemini(user_image_path, "user"),
                        self.preprocess_image_with_gemini(hairstyle_image_path, "hairstyle"),
                        return_exceptions=True
                    )
                )

                # 处理可能的异常结果
                if isinstance(processed_user_image, Exception):
                    print(f"[{thread_name}] 用户图像预处理失败: {processed_user_image}")
                    processed_user_image = user_image_path

                if isinstance(processed_hairstyle_image, Exception):
                    print(f"[{thread_name}] 发型图像预处理失败: {processed_hairstyle_image}")
                    processed_hairstyle_image = hairstyle_image_path

                print(f"[{thread_name}] 图像预处理完成")
                return processed_user_image, processed_hairstyle_image

            finally:
                loop.close()

        except Exception as e:
            print(f"[{thread_name}] 并发预处理失败: {e}")
            print(f"[{thread_name}] 使用原图继续...")
            return user_image_path, hairstyle_image_path

    def upload_image(self, image_path):
        """Upload image to RunningHub server and return fileName"""
        corrected_path = image_path
        
        conn = http.client.HTTPSConnection(self.host)
        dataList = []
        boundary = 'wL36Yn8afVp8Ag7AmP8qZ0SA4n1v9T'
        
        dataList.append(encode('--' + boundary))
        dataList.append(encode('Content-Disposition: form-data; name=apiKey;'))
        dataList.append(encode('Content-Type: {}'.format('text/plain')))
        dataList.append(encode(''))
        dataList.append(encode(self.api_key))
        
        dataList.append(encode('--' + boundary))
        filename = os.path.basename(corrected_path)
        dataList.append(encode('Content-Disposition: form-data; name=file; filename={0}'.format(filename)))
        
        fileType = mimetypes.guess_type(corrected_path)[0] or 'application/octet-stream'
        dataList.append(encode('Content-Type: {}'.format(fileType)))
        dataList.append(encode(''))
        
        with open(corrected_path, 'rb') as f:
            dataList.append(f.read())
            
        dataList.append(encode('--' + boundary))
        dataList.append(encode('Content-Disposition: form-data; name=fileType;'))
        dataList.append(encode('Content-Type: {}'.format('text/plain')))
        dataList.append(encode(''))
        dataList.append(encode("image"))
        dataList.append(encode('--'+boundary+'--'))
        dataList.append(encode(''))
        
        body = b'\r\n'.join(dataList)
        headers = {
            'Host': self.host,
            'Content-type': 'multipart/form-data; boundary={}'.format(boundary)
        }
        
        try:
            conn.request("POST", "/task/openapi/upload", body, headers)
            res = conn.getresponse()
            data = res.read()
            result = json.loads(data.decode("utf-8"))
            
            if result.get("code") == 0:
                print(f"Upload successful for {image_path}: {result['data']['fileName']}")
                return result["data"]["fileName"]
            else:
                print(f"Upload failed for {image_path}: {result}")
                print(f"API Response: {result}")
                return None
        except Exception as e:
            print(f"Error uploading {image_path}: {e}")
            return None
        finally:
            conn.close()
            # Clean up temporary corrected file if it was created
            if corrected_path != image_path and os.path.exists(corrected_path):
                try:
                    os.remove(corrected_path)
                except:
                    pass
    
    def run_hairstyle_task(self, hairstyle_filename, user_filename, max_retries=10, retry_delay=20, cancel_check_func=None):
        """Run AI hairstyle transfer task with retry mechanism for TASK_QUEUE_MAXED"""
        start_time = time.time()  # 记录开始时间
        
        payload = json.dumps({
            "webappId": self.webapp_id,
            "apiKey": self.api_key,
            "nodeInfoList": [
                {
                    "nodeId": "901",
                    "fieldName": "image",
                    "fieldValue": hairstyle_filename,
                    "description": "hair"
                },
                {
                    "nodeId": "239",
                    "fieldName": "image",
                    "fieldValue": user_filename,
                    "description": "user"
                }
            ],
        })

        headers = {
            'Host': self.host,
            'Content-Type': 'application/json'
        }

        for attempt in range(max_retries):
            # 检查是否需要取消
            if cancel_check_func and cancel_check_func():
                print(f"任务在排队阶段被取消 (attempt {attempt + 1}/{max_retries})")
                return None

            conn = http.client.HTTPSConnection(self.host)
            try:
                conn.request("POST", "/task/openapi/ai-app/run", payload, headers)
                res = conn.getresponse()
                data = res.read()
                result = json.loads(data.decode("utf-8"))

                if result.get("code") == 0:
                    end_time = time.time()  # 记录结束时间
                    elapsed_time = end_time - start_time
                    self.task_times.append(elapsed_time)
                    self.task_count += 1
                    print(f"Task started successfully: {result['data']['taskId']} (耗时: {elapsed_time:.2f}秒)")
                    return result["data"]["taskId"]
                elif result.get("msg") in ["TASK_QUEUE_MAXED", "TASK_INSTANCE_MAXED"]:
                    print(f"Task queue is full (attempt {attempt + 1}/{max_retries}), waiting {retry_delay} seconds before retry...")
                    if attempt < max_retries - 1:  # Don't sleep on the last attempt
                        # 在睡眠期间也要检查取消状态
                        for i in range(retry_delay):
                            if cancel_check_func and cancel_check_func():
                                print(f"任务在等待重试期间被取消")
                                return None
                            time.sleep(1)
                        continue
                    else: 
                        end_time = time.time()  # 记录结束时间（失败时）
                        elapsed_time = end_time - start_time
                        self.task_times.append(elapsed_time)
                        self.task_count += 1
                        print(f"Max retries reached, task queue still full (总耗时: {elapsed_time:.2f}秒)")
                        return None
                else:
                    end_time = time.time()  # 记录结束时间（失败时）
                    elapsed_time = end_time - start_time
                    self.task_times.append(elapsed_time)
                    self.task_count += 1
                    print(f"Task failed: {result} (耗时: {elapsed_time:.2f}秒)")
                    print(f"API Response: {result}")
                    return None
            except Exception as e:
                end_time = time.time()  # 记录结束时间（异常时）
                elapsed_time = end_time - start_time
                self.task_times.append(elapsed_time)
                self.task_count += 1
                print(f"Error running task (attempt {attempt + 1}/{max_retries}): {e} (耗时: {elapsed_time:.2f}秒)")
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                    continue
                else:
                    return None
            finally:
                conn.close()

        return None

    def run_color_task(self, hair_filename, user_filename, max_retries=10, retry_delay=20, cancel_check_func=None):
        """Run AI color transfer task with retry mechanism for TASK_QUEUE_MAXED"""
        if not self.color_webapp_id:
            raise ValueError("Color webapp ID is required. Set RUNNINGHUB_COLOR_WEBAPP_ID environment variable.")

        start_time = time.time()  # 记录开始时间

        payload = json.dumps({
            "webappId": self.color_webapp_id,
            "apiKey": self.api_key,
            "nodeInfoList": [
                {
                    "nodeId": "1",
                    "fieldName": "image",
                    "fieldValue": user_filename,
                    "description": "user"
                },
                {
                    "nodeId": "200",
                    "fieldName": "image",
                    "fieldValue": hair_filename,
                    "description": "hair"
                }
            ],
        })

        headers = {
            'Host': self.host,
            'Content-Type': 'application/json'
        }

        for attempt in range(max_retries):
            # 检查是否需要取消
            if cancel_check_func and cancel_check_func():
                print(f"颜色换装任务在排队阶段被取消 (attempt {attempt + 1}/{max_retries})")
                return None

            conn = http.client.HTTPSConnection(self.host)
            try:
                conn.request("POST", "/task/openapi/ai-app/run", payload, headers)
                res = conn.getresponse()
                data = res.read()
                result = json.loads(data.decode("utf-8"))

                if result.get("code") == 0:
                    end_time = time.time()  # 记录结束时间
                    elapsed_time = end_time - start_time
                    self.task_times.append(elapsed_time)
                    self.task_count += 1
                    print(f"Color task started successfully: {result['data']['taskId']} (耗时: {elapsed_time:.2f}秒)")
                    return result["data"]["taskId"]
                elif result.get("msg") in ["TASK_QUEUE_MAXED", "TASK_INSTANCE_MAXED"]:
                    print(f"Color task queue is full (attempt {attempt + 1}/{max_retries}), waiting {retry_delay} seconds before retry...")
                    if attempt < max_retries - 1:  # Don't sleep on the last attempt
                        # 在睡眠期间也要检查取消状态
                        for i in range(retry_delay):
                            if cancel_check_func and cancel_check_func():
                                print(f"颜色换装任务在等待重试期间被取消")
                                return None
                            time.sleep(1)
                        continue
                    else:
                        end_time = time.time()  # 记录结束时间（失败时）
                        elapsed_time = end_time - start_time
                        self.task_times.append(elapsed_time)
                        self.task_count += 1
                        print(f"Max retries reached, color task queue still full (总耗时: {elapsed_time:.2f}秒)")
                        return None
                else:
                    end_time = time.time()  # 记录结束时间（失败时）
                    elapsed_time = end_time - start_time
                    self.task_times.append(elapsed_time)
                    self.task_count += 1
                    print(f"Color task failed: {result} (耗时: {elapsed_time:.2f}秒)")
                    print(f"API Response: {result}")
                    return None
            except Exception as e:
                end_time = time.time()  # 记录结束时间（异常时）
                elapsed_time = end_time - start_time
                self.task_times.append(elapsed_time)
                self.task_count += 1
                print(f"Error running color task (attempt {attempt + 1}/{max_retries}): {e} (耗时: {elapsed_time:.2f}秒)")
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                    continue
                else:
                    return None
            finally:
                conn.close()

        return None
    
    def check_task_status(self, task_id):
        """Check task status"""
        conn = http.client.HTTPSConnection(self.host)
        payload = json.dumps({
            "apiKey": self.api_key,
            "taskId": task_id
        })
        
        headers = {
            'Host': self.host,
            'Content-Type': 'application/json'
        }
        
        try:
            conn.request("POST", "/task/openapi/status", payload, headers)
            res = conn.getresponse()
            data = res.read()
            result = json.loads(data.decode("utf-8"))
            
            if result.get("code") == 0:
                return result["data"]
            else:
                print(f"Status check failed: {result}")
                return None
        except Exception as e:
            print(f"Error checking status: {e}")
            return None
        finally:
            conn.close()
    
    def get_task_results(self, task_id):
        """Get task results"""
        conn = http.client.HTTPSConnection(self.host)
        payload = json.dumps({
            "apiKey": self.api_key,
            "taskId": task_id
        })

        headers = {
            'Host': self.host,
            'Content-Type': 'application/json'
        }

        try:
            conn.request("POST", "/task/openapi/outputs", payload, headers)
            res = conn.getresponse()
            data = res.read()
            result = json.loads(data.decode("utf-8"))

            if result.get("code") == 0:
                return result["data"]
            else:
                print(f"Get results failed: {result}")
                return None
        except Exception as e:
            print(f"Error getting results: {e}")
            return None
        finally:
            conn.close()

    def cancel_task(self, task_id):
        """Cancel task"""
        conn = http.client.HTTPSConnection(self.host)
        payload = json.dumps({
            "apiKey": self.api_key,
            "taskId": task_id
        })

        headers = {
            'Host': self.host,
            'Content-Type': 'application/json'
        }

        try:
            conn.request("POST", "/task/openapi/cancel", payload, headers)
            res = conn.getresponse()
            data = res.read()
            result = json.loads(data.decode("utf-8"))

            if result.get("code") == 0:
                print(f"Task cancelled successfully: {task_id}")
                return True
            else:
                print(f"Cancel task failed: {result}")
                return False
        except Exception as e:
            print(f"Error cancelling task: {e}")
            return False
        finally:
            conn.close()
    
    def download_image(self, url, save_path):
        """Download image from URL"""
        try:
            response = requests.get(url)
            if response.status_code == 200:
                with open(save_path, 'wb') as f:
                    f.write(response.content)
                return True
            else:
                print(f"Failed to download {url}")
                return False
        except Exception as e:
            print(f"Error downloading {url}: {e}")
            return False
    
    def create_combined_image(self, hairstyle_path, user_path, result_paths, output_path):
        """Create a combined image with hairstyle reference, user photo, and all generated results"""
        try:
            # Open hairstyle and user images
            hairstyle_img = Image.open(hairstyle_path)
            user_img = Image.open(user_path)
            
            # Open all result images
            result_imgs = []
            for result_path in result_paths:
                if os.path.exists(result_path):
                    result_imgs.append(Image.open(result_path))
            
            if not result_imgs:
                print("No result images found")
                return False
            
            # Collect all images
            all_imgs = [hairstyle_img, user_img] + result_imgs
            
            # Convert to RGB if necessary
            for i, img in enumerate(all_imgs):
                if img.mode != 'RGB':
                    all_imgs[i] = img.convert('RGB')
            
            # Define target height (use the minimum height among all images, but at least 512px)
            target_height = max(512, min(img.height for img in all_imgs))
            
            # Resize all images to the same height while maintaining aspect ratio
            def resize_to_height(img, target_height):
                aspect_ratio = img.width / img.height
                target_width = int(target_height * aspect_ratio)
                return img.resize((target_width, target_height), Image.Resampling.LANCZOS)
            
            resized_imgs = [resize_to_height(img, target_height) for img in all_imgs]
            
            # Calculate total width
            total_width = sum(img.width for img in resized_imgs)
            
            # Create new image for the combined result
            combined_img = Image.new('RGB', (total_width, target_height), (255, 255, 255))
            
            # Paste images side by side
            x_offset = 0
            for img in resized_imgs:
                combined_img.paste(img, (x_offset, 0))
                x_offset += img.width
            
            # Save the combined image
            combined_img.save(output_path, 'PNG', quality=95)
            print(f"Combined image saved: {output_path}")
            return True
            
        except Exception as e:
            print(f"Error creating combined image: {e}")
            return False
    
    def resize_image_for_word(self, image_path, max_width=2.5):
        """Resize image to fit in Word document"""
        try:
            with Image.open(image_path) as img:
                width, height = img.size
                aspect_ratio = height / width
                
                if width > max_width * 96:  # 96 DPI default
                    new_width = max_width
                    new_height = new_width * aspect_ratio
                    return new_width, new_height
                else:
                    return width / 96, height / 96
        except:
            return max_width, max_width
    
    def process_single_combination_with_timeout(self, task_info):
        """Process a single user-hairstyle combination with timeout control"""
        start_time = time.time()
        thread_name = threading.current_thread().name
        user_file = task_info[2]
        hairstyle_file = task_info[3]

        try:
            print(f"[{thread_name}] 开始处理任务 (超时限制: {self.task_timeout}秒): {user_file} + {hairstyle_file}")
            result = self.process_single_combination(task_info)
            end_time = time.time()
            elapsed = end_time - start_time

            # 检查是否超时
            if elapsed > self.task_timeout:
                self.timeout_count += 1
                print(f"[{thread_name}] ⚠️ 任务超时 (耗时: {elapsed:.2f}秒): {user_file} + {hairstyle_file}")
                return None

            print(f"[{thread_name}] 任务完成，耗时: {elapsed:.2f}秒: {user_file} + {hairstyle_file}")
            return result

        except Exception as e:
            end_time = time.time()
            elapsed = end_time - start_time
            print(f"[{thread_name}] ❌ 任务异常 (耗时: {elapsed:.2f}秒): {user_file} + {hairstyle_file}")
            print(f"[{thread_name}] 异常详情: {e}")
            return None

    def process_single_combination(self, task_info):
        """Process a single user-hairstyle combination with Gemini preprocessing"""
        user_full_path, hairstyle_full_path, user_file, hairstyle_file, gender_name, results_dir = task_info

        print(f"[{threading.current_thread().name}] Processing: {user_file} + {hairstyle_file}")

        try:
            # Step 1: Gemini预处理图像
            print(f"[{threading.current_thread().name}] Step 1: Gemini preprocessing...")
            processed_user_path, processed_hairstyle_path = self.preprocess_images_concurrently(
                user_full_path, hairstyle_full_path
            )

            # Step 2: Upload processed images
            print(f"[{threading.current_thread().name}] Step 2: Uploading processed images...")
            user_filename = self.upload_image(processed_user_path)
            if not user_filename:
                print(f"[{threading.current_thread().name}] Failed to upload user image, trying original...")
                user_filename = self.upload_image(user_full_path)
                if not user_filename:
                    return

            hairstyle_filename = self.upload_image(processed_hairstyle_path)
            if not hairstyle_filename:
                print(f"[{threading.current_thread().name}] Failed to upload hairstyle image, trying original...")
                hairstyle_filename = self.upload_image(hairstyle_full_path)
                if not hairstyle_filename:
                    return
            
            # Run task
            print(f"[{threading.current_thread().name}] Running hairstyle transfer task...")
            task_id = self.run_hairstyle_task(hairstyle_filename, user_filename)
            if not task_id:
                return
            
            # Wait for completion
            print(f"[{threading.current_thread().name}] Task {task_id} started, waiting for completion...")
            max_wait = 1000  # 5 minutes max
            wait_time = 0
            
            while wait_time < max_wait:
                status = self.check_task_status(task_id)
                if status == "SUCCESS":
                    break
                elif status in ["FAILED", "CANCELLED"]:
                    print(f"[{threading.current_thread().name}] Task failed with status: {status}")
                    return
                
                time.sleep(10)
                wait_time += 10
                if wait_time % 10 == 0:  # Print every 10 seconds
                    print(f"[{threading.current_thread().name}] Still processing... ({wait_time}s)")
            
            if status != "SUCCESS":
                print(f"[{threading.current_thread().name}] Task did not complete successfully: {status}")
                return
            
            # Get results
            print(f"[{threading.current_thread().name}] Getting results...")
            results = self.get_task_results(task_id)
            if not results:
                return
            
            # Download result images and create combined images
            result_paths = []
            result_filenames = []
            
            # Download all result images first
            for i, result in enumerate(results):
                result_url = result.get("fileUrl")
                if result_url:
                    result_filename = f"{gender_name}_{user_file}_{hairstyle_file}_result_{i}.png"
                    result_path = os.path.join(results_dir, result_filename)
                    
                    if self.download_image(result_url, result_path):
                        result_paths.append(result_path)
                        result_filenames.append(result_filename)
            
            # Create one combined image with all results (original hairstyle + original user + results)
            if result_paths:
                combined_filename = f"{gender_name}_{user_file}_{hairstyle_file}_combined_all.png"
                combined_path = os.path.join(results_dir, combined_filename)

                # Use original images for the combined image to show the transformation
                if self.create_combined_image(hairstyle_full_path, user_full_path, result_paths, combined_path):
                    print(f"[{threading.current_thread().name}] Created combined image: {combined_filename}")

                # Store result info (thread-safe) - include both original and processed paths
                with self.results_lock:
                    self.results.append({
                        'gender': gender_name,
                        'user_image': user_full_path,  # 保留原始路径用于记录
                        'hairstyle_image': hairstyle_full_path,  # 保留原始路径用于记录
                        'processed_user_image': processed_user_path,  # 新增预处理路径
                        'processed_hairstyle_image': processed_hairstyle_path,  # 新增预处理路径
                        'result_images': result_paths,
                        'combined_image': combined_path if os.path.exists(combined_path) else None,
                        'user_filename': user_file,
                        'hairstyle_filename': hairstyle_file,
                        'result_filenames': result_filenames,
                        'combined_filename': combined_filename
                    })
            
            print(f"[{threading.current_thread().name}] Completed: {user_file} + {hairstyle_file}")
            
        except Exception as e:
            print(f"[{threading.current_thread().name}] Error processing {user_file} + {hairstyle_file}: {e}")
    
    def process_gender_folder(self, gender_path, gender_name):
        """Process all combinations for a gender (man/woman) with concurrent processing"""
        hairstyle_path = os.path.join(gender_path, "hairstyle")
        user_path = os.path.join(gender_path, "user")
        
        if not os.path.exists(hairstyle_path) or not os.path.exists(user_path):
            print(f"Missing hairstyle or user folder for {gender_name}")
            return
        
        hairstyle_files = [f for f in os.listdir(hairstyle_path) if f.lower().endswith(('.jpg', '.jpeg', '.png','.JPG', '.JPEG', '.PNG'))]
        user_files = [f for f in os.listdir(user_path) if f.lower().endswith(('.jpg', '.jpeg', '.png','.JPG', '.JPEG', '.PNG'))]
        
        # # For women, randomly select 50 hairstyles
        # if gender_name == "woman" and len(hairstyle_files) > 50:
        #     hairstyle_files = random.sample(hairstyle_files, 50)
        #     print(f"Randomly selected 50 hairstyles for women from {len(os.listdir(hairstyle_path))} total")
        
        print(f"Processing {gender_name}: {len(hairstyle_files)} hairstyles × {len(user_files)} users = {len(hairstyle_files) * len(user_files)} combinations")
        
        results_dir = os.path.join(self.data_dir, f"results_{gender_name}_{datetime.now().strftime('%m%d')}_")
        os.makedirs(results_dir, exist_ok=True)
        
        # Create task list
        tasks = []
        for user_file in user_files:
            for hairstyle_file in hairstyle_files:
                user_full_path = os.path.join(user_path, user_file)
                hairstyle_full_path = os.path.join(hairstyle_path, hairstyle_file)
                task_info = (user_full_path, hairstyle_full_path, user_file, hairstyle_file, gender_name, results_dir)
                tasks.append(task_info)
                # break
        
        # Process tasks concurrently
        print(f"Starting concurrent processing with {self.max_workers} workers (timeout: {self.task_timeout}s per task)...")
        with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all tasks
            future_to_task = {executor.submit(self.process_single_combination_with_timeout, task): task for task in tasks}

            # Process completed tasks
            completed = 0
            successful = 0
            failed = 0
            timeout_tasks = 0

            for future in concurrent.futures.as_completed(future_to_task, timeout=None):
                completed += 1
                task = future_to_task[future]
                user_file, hairstyle_file = task[2], task[3]

                try:
                    # 使用任务级别的超时
                    result = future.result(timeout=self.task_timeout + 30)  # 给额外30秒的缓冲时间
                    if result is not None:
                        successful += 1
                        print(f"✅ Progress: {completed}/{len(tasks)} - Success: {successful}, Failed: {failed}, Timeout: {timeout_tasks}")
                    else:
                        failed += 1
                        print(f"❌ Progress: {completed}/{len(tasks)} - Success: {successful}, Failed: {failed}, Timeout: {timeout_tasks}")

                except concurrent.futures.TimeoutError:
                    timeout_tasks += 1
                    failed += 1
                    print(f"⏰ Future timeout: {user_file} + {hairstyle_file}")
                    print(f"⚠️ Progress: {completed}/{len(tasks)} - Success: {successful}, Failed: {failed}, Timeout: {timeout_tasks}")

                except Exception as exc:
                    failed += 1
                    print(f"💥 Task {user_file} + {hairstyle_file} generated an exception: {exc}")
                    print(f"❌ Progress: {completed}/{len(tasks)} - Success: {successful}, Failed: {failed}, Timeout: {timeout_tasks}")
            
            print(f"\n=== 处理完成统计 ===")
            print(f"总任务数: {len(tasks)}")
            print(f"成功完成: {successful}")
            print(f"失败任务: {failed}")
            print(f"超时任务: {self.timeout_count}")
            print(f"成功率: {(successful/len(tasks)*100):.1f}%")
            print(f"===================")
        
        print(f"Completed processing {gender_name} folder")
    
    def create_word_document(self, output_path="hairstyle_results.docx"):
        """Create Word document with all results"""
        doc = Document()
        doc.add_heading('发型换装结果', 0)
        
        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'总共处理: {len(self.results)} 个组合')
        
        for i, result in enumerate(self.results):
            doc.add_heading(f'结果 {i+1}: {result["gender"]} - {result["user_filename"]} + {result["hairstyle_filename"]}', level=1)
            
            # Add combined image if available
            if result.get('combined_image') and os.path.exists(result['combined_image']):
                doc.add_paragraph('拼接图片 (发型参考 + 用户照片 + 生成结果):')
                width, height = self.resize_image_for_word(result['combined_image'], max_width=6.0)  # Wider for combined image
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(result['combined_image'], width=Inches(width), height=Inches(height))
                doc.add_paragraph()  # Add some space
            
            # Create table for individual images
            doc.add_paragraph('单独图片:')
            result_images = result.get('result_images', [])
            num_cols = 2 + len(result_images)  # hairstyle + user + result images
            table = doc.add_table(rows=2, cols=num_cols)
            table.style = 'Table Grid'
            
            # Headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '发型参考图'
            hdr_cells[1].text = '用户照片'
            for j in range(len(result_images)):
                hdr_cells[2 + j].text = f'生成结果{j+1}'
            
            # Images
            img_cells = table.rows[1].cells
            
            # Add hairstyle image
            if os.path.exists(result['hairstyle_image']):
                width, height = self.resize_image_for_word(result['hairstyle_image'])
                paragraph = img_cells[0].paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(result['hairstyle_image'], width=Inches(width), height=Inches(height))
            
            # Add user image
            if os.path.exists(result['user_image']):
                width, height = self.resize_image_for_word(result['user_image'])
                paragraph = img_cells[1].paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(result['user_image'], width=Inches(width), height=Inches(height))
            
            # Add result images
            for j, result_image in enumerate(result_images):
                if os.path.exists(result_image):
                    width, height = self.resize_image_for_word(result_image)
                    paragraph = img_cells[2 + j].paragraphs[0]
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(result_image, width=Inches(width), height=Inches(height))
            
            doc.add_page_break()
        
        doc.save(output_path)
        print(f"Word document saved: {output_path}")

    def get_cache_info(self):
        """获取缓存信息"""
        cache_info = {
            'user': {'total_files': 0, 'total_size': 0, 'files': []},
            'hairstyle': {'total_files': 0, 'total_size': 0, 'files': []}
        }

        for image_type in ['user', 'hairstyle']:
            cache_dir = os.path.join(self.data_dir, f"gemini_processed_{image_type}")
            if os.path.exists(cache_dir):
                try:
                    for filename in os.listdir(cache_dir):
                        if filename == 'cache_index.json':
                            continue
                        filepath = os.path.join(cache_dir, filename)
                        if os.path.isfile(filepath):
                            file_stat = os.stat(filepath)
                            cache_info[image_type]['files'].append({
                                'filename': filename,
                                'filepath': filepath,
                                'size': file_stat.st_size,
                                'modified_time': file_stat.st_mtime,
                                'created_time': file_stat.st_ctime
                            })
                            cache_info[image_type]['total_size'] += file_stat.st_size
                            cache_info[image_type]['total_files'] += 1
                except Exception as e:
                    print(f"获取{image_type}缓存信息失败: {e}")

        return cache_info

    def clean_old_cache(self, max_age_hours=24, max_total_size_mb=100):
        """清理旧的缓存文件"""
        current_time = time.time()
        max_age_seconds = max_age_hours * 3600
        max_total_size_bytes = max_total_size_mb * 1024 * 1024

        total_cleaned_files = 0
        total_cleaned_size = 0

        for image_type in ['user', 'hairstyle']:
            cache_dir = os.path.join(self.data_dir, f"gemini_processed_{image_type}")
            if not os.path.exists(cache_dir):
                continue

            try:
                cache_index_path = os.path.join(cache_dir, "cache_index.json")
                cache_index = {}

                # 读取缓存索引
                if os.path.exists(cache_index_path):
                    try:
                        with open(cache_index_path, 'r', encoding='utf-8') as f:
                            cache_index = json.load(f)
                    except:
                        cache_index = {}

                # 获取所有缓存文件信息
                cache_files = []
                for filename in os.listdir(cache_dir):
                    if filename == 'cache_index.json':
                        continue
                    filepath = os.path.join(cache_dir, filename)
                    if os.path.isfile(filepath):
                        file_stat = os.stat(filepath)
                        cache_files.append({
                            'filename': filename,
                            'filepath': filepath,
                            'size': file_stat.st_size,
                            'modified_time': file_stat.st_mtime
                        })

                # 按修改时间排序（旧的在前）
                cache_files.sort(key=lambda x: x['modified_time'])

                cleaned_files_in_type = 0
                cleaned_size_in_type = 0

                # 计算总大小
                total_size = sum(f['size'] for f in cache_files)

                # 清理策略1: 删除超过指定时间的文件
                files_to_remove = []
                for file_info in cache_files:
                    file_age = current_time - file_info['modified_time']
                    if file_age > max_age_seconds:
                        files_to_remove.append(file_info)

                # 清理策略2: 如果总大小超过限制，删除最旧的文件
                if total_size > max_total_size_bytes:
                    remaining_files = [f for f in cache_files if f not in files_to_remove]
                    remaining_size = sum(f['size'] for f in remaining_files)

                    for file_info in remaining_files:
                        if remaining_size <= max_total_size_bytes:
                            break
                        files_to_remove.append(file_info)
                        remaining_size -= file_info['size']

                # 执行删除操作
                for file_info in files_to_remove:
                    try:
                        os.remove(file_info['filepath'])
                        cleaned_files_in_type += 1
                        cleaned_size_in_type += file_info['size']

                        # 从缓存索引中移除对应条目
                        filename_hash = None
                        for hash_key, index_info in cache_index.items():
                            if index_info.get('processed_path') == file_info['filepath']:
                                filename_hash = hash_key
                                break

                        if filename_hash:
                            del cache_index[filename_hash]

                        print(f"删除缓存文件: {file_info['filename']} ({file_info['size'] / 1024:.1f}KB)")

                    except Exception as e:
                        print(f"删除文件失败 {file_info['filepath']}: {e}")

                # 更新缓存索引
                if cleaned_files_in_type > 0:
                    try:
                        with open(cache_index_path, 'w', encoding='utf-8') as f:
                            json.dump(cache_index, f, ensure_ascii=False, indent=2)
                    except Exception as e:
                        print(f"更新{image_type}缓存索引失败: {e}")

                total_cleaned_files += cleaned_files_in_type
                total_cleaned_size += cleaned_size_in_type

                if cleaned_files_in_type > 0:
                    print(f"清理{image_type}缓存: {cleaned_files_in_type}个文件, {cleaned_size_in_type / 1024:.1f}KB")

            except Exception as e:
                print(f"清理{image_type}缓存目录失败: {e}")

        if total_cleaned_files > 0:
            print(f"缓存清理完成: 总计删除{total_cleaned_files}个文件, {total_cleaned_size / 1024:.1f}KB")
        else:
            print("无需清理缓存文件")

        return {
            'cleaned_files': total_cleaned_files,
            'cleaned_size': total_cleaned_size
        }

    def get_disk_usage(self):
        """获取磁盘使用情况"""
        try:
            import shutil
            total, used, free = shutil.disk_usage(self.data_dir)
            return {
                'total': total,
                'used': used,
                'free': free,
                'usage_percent': (used / total) * 100
            }
        except Exception as e:
            print(f"获取磁盘使用情况失败: {e}")
            return None

    def delete_cache_file(self, file_path, image_type):
        """删除指定的缓存文件"""
        try:
            # 验证文件路径是否在缓存目录内（安全检查）
            cache_dir = os.path.join(self.data_dir, f"gemini_processed_{image_type}")
            normalized_file_path = os.path.normpath(file_path)
            normalized_cache_dir = os.path.normpath(cache_dir)

            if not normalized_file_path.startswith(normalized_cache_dir):
                print(f"安全检查失败: 文件路径不在缓存目录内 {file_path}")
                return False

            # 检查文件是否存在
            if not os.path.exists(file_path):
                print(f"文件不存在: {file_path}")
                return False

            # 获取文件大小（用于统计）
            file_size = os.path.getsize(file_path)

            # 删除文件
            os.remove(file_path)

            # 从缓存索引中移除对应条目
            cache_index_path = os.path.join(cache_dir, "cache_index.json")
            if os.path.exists(cache_index_path):
                try:
                    with open(cache_index_path, 'r', encoding='utf-8') as f:
                        cache_index = json.load(f)

                    # 查找并删除对应的索引条目
                    hash_to_remove = None
                    for hash_key, index_info in cache_index.items():
                        if index_info.get('processed_path') == file_path:
                            hash_to_remove = hash_key
                            break

                    if hash_to_remove:
                        del cache_index[hash_to_remove]

                        # 更新索引文件
                        with open(cache_index_path, 'w', encoding='utf-8') as f:
                            json.dump(cache_index, f, ensure_ascii=False, indent=2)

                except Exception as e:
                    print(f"更新缓存索引失败: {e}")

            print(f"删除缓存文件成功: {os.path.basename(file_path)} ({file_size / 1024:.1f}KB)")
            return True

        except Exception as e:
            print(f"删除缓存文件失败 {file_path}: {e}")
            return False

    def get_cache_files_detailed(self):
        """获取详细的缓存文件列表"""
        cache_files = {
            'user': [],
            'hairstyle': []
        }

        for image_type in ['user', 'hairstyle']:
            cache_dir = os.path.join(self.data_dir, f"gemini_processed_{image_type}")
            if os.path.exists(cache_dir):
                try:
                    # 读取缓存索引
                    cache_index_path = os.path.join(cache_dir, "cache_index.json")
                    cache_index = {}
                    if os.path.exists(cache_index_path):
                        try:
                            with open(cache_index_path, 'r', encoding='utf-8') as f:
                                cache_index = json.load(f)
                        except:
                            cache_index = {}

                    # 获取所有缓存文件
                    for filename in os.listdir(cache_dir):
                        if filename == 'cache_index.json':
                            continue

                        filepath = os.path.join(cache_dir, filename)
                        if os.path.isfile(filepath):
                            try:
                                file_stat = os.stat(filepath)

                                # 查找对应的原始文件信息
                                original_filename = None
                                original_path = None
                                for hash_key, index_info in cache_index.items():
                                    if index_info.get('processed_path') == filepath:
                                        original_filename = index_info.get('original_filename', 'Unknown')
                                        original_path = index_info.get('original_path', '')
                                        break

                                cache_files[image_type].append({
                                    'filename': filename,
                                    'filepath': filepath,
                                    'original_filename': original_filename or filename,
                                    'original_path': original_path or '',
                                    'size': file_stat.st_size,
                                    'size_mb': file_stat.st_size / (1024 * 1024),
                                    'modified_time': file_stat.st_mtime,
                                    'created_time': file_stat.st_ctime,
                                    'modified_time_str': datetime.fromtimestamp(file_stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                                    'created_time_str': datetime.fromtimestamp(file_stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
                                })
                            except Exception as e:
                                print(f"获取文件 {filename} 信息失败: {e}")

                    # 按修改时间排序（新的在前）
                    cache_files[image_type].sort(key=lambda x: x['modified_time'], reverse=True)

                except Exception as e:
                    print(f"获取{image_type}缓存文件详情失败: {e}")

        return cache_files

    def get_average_task_time(self):
        """计算并显示run_hairstyle_task和Gemini预处理的统计信息"""

        # RunningHub任务统计
        if not self.task_times:
            print("没有RunningHub任务运行记录")
            runninghub_avg = 0.0
        else:
            total_time = sum(self.task_times)
            runninghub_avg = total_time / len(self.task_times)
            min_time = min(self.task_times)
            max_time = max(self.task_times)

            print(f"\n=== RunningHub任务统计 ===")
            print(f"总任务数: {len(self.task_times)}")
            print(f"总运行时间: {total_time:.2f}秒")
            print(f"平均运行时间: {runninghub_avg:.2f}秒")
            print(f"最短运行时间: {min_time:.2f}秒")
            print(f"最长运行时间: {max_time:.2f}秒")
            print(f"========================\n")

        # Gemini预处理统计
        if not self.gemini_times:
            print("没有Gemini预处理记录")
            gemini_avg = 0.0
        else:
            total_gemini_time = sum(self.gemini_times)
            gemini_avg = total_gemini_time / len(self.gemini_times)
            min_gemini_time = min(self.gemini_times)
            max_gemini_time = max(self.gemini_times)

            print(f"=== Gemini预处理统计 ===")
            print(f"总预处理请求数: {len(self.gemini_times)}")
            print(f"成功处理数: {self.gemini_success_count}")
            print(f"失败处理数: {self.gemini_fail_count}")
            print(f"成功率: {(self.gemini_success_count / (self.gemini_success_count + self.gemini_fail_count) * 100):.1f}%" if (self.gemini_success_count + self.gemini_fail_count) > 0 else "N/A")
            print(f"总预处理时间: {total_gemini_time:.2f}秒")
            print(f"平均预处理时间: {gemini_avg:.2f}秒")
            print(f"最短预处理时间: {min_gemini_time:.2f}秒")
            print(f"最长预处理时间: {max_gemini_time:.2f}秒")
            print(f"========================\n")

        # 综合统计
        total_processed_combinations = len(self.results)
        if total_processed_combinations > 0 or self.timeout_count > 0:
            print(f"=== 综合处理统计 ===")
            print(f"处理的图像组合数: {total_processed_combinations}")
            print(f"超时任务数: {self.timeout_count}")
            if self.timeout_count > 0:
                total_attempts = total_processed_combinations + self.timeout_count
                print(f"任务成功率: {(total_processed_combinations/total_attempts*100):.1f}%")
                print(f"任务超时率: {(self.timeout_count/total_attempts*100):.1f}%")
            print(f"平均RunningHub任务时间: {runninghub_avg:.2f}秒")
            print(f"平均Gemini预处理时间: {gemini_avg:.2f}秒")
            print(f"任务超时限制: {self.task_timeout}秒")
            print(f"===================\n")

        return runninghub_avg

def main():
    hair_base_path = "/Users/alex_wu/work/hair"
    
    # Set random seed for reproducible results
    random.seed(42)
    
    # 创建处理器，设置超时时间为30分钟（1800秒）
    processor = HairstyleProcessor(max_workers=2, task_timeout=600)
    
    # Process women's hairstyles (with random selection of 50)
    woman_path = os.path.join(hair_base_path, "woman")
    if os.path.exists(woman_path):
        print("Starting women's hairstyle processing...")
        processor.process_gender_folder(woman_path, "woman")
    
    # Process men's hairstyles
    man_path = os.path.join(hair_base_path, "man")
    if os.path.exists(man_path):
        print("Starting men's hairstyle processing...")
        processor.process_gender_folder(man_path, "man")

    # Create Word document with all results
    # if processor.results:
    #     processor.create_word_document("hairstyle_results.docx")
    #     print(f"Processing complete! Generated {len(processor.results)} results.")
    # else:
    #     print("No results generated.")
    
    # 显示任务运行时间统计
    processor.get_average_task_time()

if __name__ == "__main__":
    main()